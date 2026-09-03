# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_req(p, num, text, is_open=False):
    run_num = p.add_run(num + " ")
    run_num.bold = True
    run_text = p.add_run(text)
    if is_open:
        run_num.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
        run_text.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
        run_text.italic = True

# ── Trang bìa ──
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("TÀI LIỆU YÊU CẦU NGHIỆP VỤ")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Business Requirements Document (BRD)").font.size = Pt(13)

doc.add_paragraph()
info_table = doc.add_table(rows=5, cols=2)
info_data = [
    ("Dự án", "ERP CloudAZ — Module Tính Cước & Đối Soát (Billing & Dispute)"),
    ("Khách hàng", "CloudAZ / Cloudino"),
    ("Ngày", "2026-08-20"),
    ("Phiên bản", "2.0"),
    ("Tác giả", "BA Team (AI-assisted)"),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    set_cell_shading(info_table.rows[i].cells[0], "E8F0FE")
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            p.style.font.size = Pt(11)

doc.add_page_break()

# ── Lịch sử Tài liệu ──
doc.add_heading("Lịch sử Tài liệu", level=1)
hist_table = doc.add_table(rows=2, cols=4)
hist_table.style = 'Light Shading Accent 1'
headers = ["Phiên bản", "Ngày", "Tác giả", "Mô tả"]
for i, h in enumerate(headers):
    hist_table.rows[0].cells[i].text = h
hist_table.rows[1].cells[0].text = "2.0"
hist_table.rows[1].cells[1].text = "2026-08-20"
hist_table.rows[1].cells[2].text = "BA Team (AI-assisted)"
hist_table.rows[1].cells[3].text = "Cập nhật BRD từ phân tích code hiện tại (GCP, GMP, GWS Flex)"

# ── 1. Vấn đề Hiện tại ──
doc.add_heading("1. Vấn đề Hiện tại", level=1)
doc.add_paragraph(
    "CloudAZ/Cloudino là đối tác (reseller) của các hãng Cloud lớn: Google Cloud Platform (GCP), "
    "Google Marketing Platform (GMP), Google Workspace (GWS), AWS, DigitalOcean. Hàng tháng, Kế toán "
    "doanh thu phải thực hiện quy trình tính cước (billing) cho khoảng 70–80 khách hàng GCP, ~40 khách "
    "GMP và hàng chục khách GWS, sau đó gửi bảng đối soát chi phí cho khách xác nhận trước khi xuất hóa đơn."
)
doc.add_paragraph(
    "Hiện tại, toàn bộ quy trình được thực hiện thủ công trên nhiều công cụ rời rạc: Console các hãng, "
    "bảng tính Excel, hệ thống nội bộ CM (chỉ hỗ trợ Google, không hỗ trợ AWS/DO), và email. "
    "Các điểm đau chính:"
)
pain_points = [
    "Thao tác thủ công quá nhiều: copy-paste từng khách (~70-80 khách GCP × 5-6 bước), mất 1–1,5 ngày/tháng.",
    "Hệ thống CM cũ không đáp ứng: không hỗ trợ Gemini API (dịch vụ AI mới), lỗi làm tròn (rounding) thường xuyên lệch 1-2 đồng, không hỗ trợ AWS & DigitalOcean.",
    "Không có cảnh báo tự động khi: hợp đồng thay đổi, credit promotion mới, khách thêm project mới.",
    "Phụ thuộc 'trực giác kế toán': không có validation tự động, rủi ro cao khi kế toán nghỉ/thay người.",
    "Dữ liệu lưu trữ phân tán trên PC cá nhân, email, Drive — khó tra cứu lịch sử bill.",
    "Chưa có SLA xử lý tranh chấp cước, chưa có audit trail khi sửa tay bill.",
]
for pp in pain_points:
    doc.add_paragraph(pp, style='List Bullet')

# ── 2. Giải pháp Đề xuất ──
doc.add_heading("2. Giải pháp Đề xuất", level=1)
doc.add_paragraph(
    "Xây dựng module Tính Cước & Đối Soát (Billing & Dispute) trong hệ thống ERP CloudAZ mới nhằm tự động hóa "
    "toàn bộ quy trình từ khi nhận invoice hãng đến khi xuất hóa đơn cho khách. Hệ thống sẽ:"
)
solution_points = [
    "Tự động kết nối và lấy dữ liệu cước từ Console/API các hãng (GCP, GMP, GWS, mở rộng sang AWS, DO).",
    "Tự động tách và nhận diện Gemini API usage, Credit Promotion theo từng khách.",
    "Áp dụng công thức tính cước theo hợp đồng (discount, thuế nhà thầu, phí dịch vụ) tự động — cấu hình riêng cho từng khách/hợp đồng.",
    "Sinh bảng đối soát chi phí tự động dạng XLSX (template-based), hỗ trợ đối soát với bảng tính tay của kế toán.",
    "Gửi email bảng đối soát cho khách, theo dõi xác nhận, tự động nhắc nhở khi quá hạn.",
    "Lưu trữ tập trung toàn bộ lịch sử bill trên S3, hỗ trợ tra cứu theo khách/tháng/dịch vụ.",
    "Rút ngắn thời gian tính bill từ 1,5 ngày xuống mục tiêu 1 ngày hoặc ít hơn.",
]
for sp in solution_points:
    doc.add_paragraph(sp, style='List Bullet')

# ── 3. Hệ thống bị Ảnh hưởng ──
doc.add_heading("3. Hệ thống bị Ảnh hưởng", level=1)
systems = [
    "Hệ thống CM (phần mềm nội bộ hiện tại) — cần đánh giá tích hợp hoặc thay thế.",
    "Google Cloud Console (GCP, GMP, GWS) — nguồn lấy dữ liệu cước.",
    "Google Billing API / BigQuery Export — nguồn dữ liệu tự động.",
    "AWS Billing Console / Cost Explorer (mở rộng tương lai).",
    "DigitalOcean Billing API (mở rộng tương lai).",
    "Hệ thống email nội bộ — gửi bảng đối soát và theo dõi phản hồi.",
    "Hệ thống kế toán/xuất hóa đơn — nhận đầu ra từ module Billing.",
    "Hệ thống quản lý hợp đồng (SSCC) — nguồn thông tin công thức tính, discount, phụ lục.",
    "Ngân hàng Techcombank (hoặc ngân hàng khác theo HĐ) — nguồn tỷ giá.",
    "AWS S3 — lưu trữ template XLSX và file output (đã implement).",
]
for s in systems:
    doc.add_paragraph(s, style='List Bullet')

# ── 4. Giả định / Phụ thuộc ──
doc.add_heading("4. Giả định / Phụ thuộc", level=1)
assumptions = [
    "Giai đoạn 1 tập trung vào 3 dịch vụ Google (GCP, GMP, GWS Flex). AWS và DigitalOcean sẽ mở rộng ở các giai đoạn sau.",
    "ERP sẽ thay thế hoặc tích hợp với hệ thống CM hiện tại — cần làm việc với admin CM để hiểu rõ cấu trúc dữ liệu trước.",
    "Kế toán doanh thu chấp nhận quy trình đối soát mới: ERP gen số → Kế toán đối chiếu → Nếu khớp thì dùng.",
    "Thông tin hợp đồng (discount, thuế, phí dịch vụ) sẽ được nhập cấu hình trên ERP theo từng khách hàng/hợp đồng.",
    "Tỷ giá mặc định: tỷ giá bán chuyển khoản Techcombank, ngày cuối tháng chi phí. Hợp đồng đặc biệt có thể cấu hình ngân hàng và thời điểm lấy tỷ giá riêng.",
    "GWS Committed (license trả trước theo năm) không nằm trong scope billing hàng tháng của giai đoạn này.",
    "Credit promotion cần xác nhận với Sale/Ban Giám Đốc trước khi áp dụng — luồng phê duyệt cần thiết kế.",
    "Đội dev ERP được cấp quyền truy cập Console Google để phát triển và kiểm thử.",
    "Template XLSX bảng đối soát được lưu trên S3 và quản lý qua costTableTemplate model (đã implement).",
    "Dữ liệu Excel GCP/GMP cần 2 sheet: Project number + Subaccount ID (đã implement validate).",
    "GWS Flex dùng SKU pricing hardcode cho các gói Workspace (đã implement trong code).",
]
for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

# ── 5. Yêu cầu Nghiệp vụ ──
doc.add_heading("5. Yêu cầu Nghiệp vụ", level=1)

requirements = [
    ("5.1 Lấy dữ liệu cước từ hãng (Data Ingestion)", [
        ("5.1.1", "Hệ thống tự động kết nối và lấy dữ liệu cước hàng tháng từ Google Cloud Console cho tất cả khách hàng GCP (~70-80 khách, ~94 billing ID, ~600+ project).", False),
        ("5.1.2", "Hệ thống tự động tải và xử lý file CSV tổng từ Google Workspace Console cho tất cả domain GWS Flex.", False),
        ("5.1.3", "Hệ thống tự động lấy dữ liệu cước GMP từ các billing link (1 link có thể chứa tới 23 project/khách hàng khác nhau).", False),
        ("5.1.4", "Hệ thống lưu trữ 2 dạng dữ liệu upload song song: theo Billing ID (~94 dòng) và theo Project (~600+ dòng).", False),
        ("5.1.5", "Hệ thống tự động lọc bỏ dòng GWS Committed khỏi dữ liệu Flex khi domain có cả 2 gói.", False),
        ("5.1.6", "[CẦN XÁC NHẬN] Hệ thống cho phép nhập tỷ giá USD-VND (mặc định: tỷ giá bán chuyển khoản Techcombank, ngày cuối tháng). Hợp đồng đặc biệt có thể cấu hình ngân hàng và thời điểm lấy tỷ giá riêng.", True),
    ]),
    ("5.2 Tách & Nhận diện Gemini API và Credit Promotion", [
        ("5.2.1", "Hệ thống tự động phát hiện và tách riêng lượng dùng Gemini API cho từng khách GCP (group by Service trên Console).", False),
        ("5.2.2", "Nếu lượng Gemini API quá nhỏ (ngưỡng linh hoạt, mặc định < $0.07), hệ thống gộp vào tổng usage thay vì tách riêng.", False),
        ("5.2.3", "Hệ thống tự động phát hiện Credit Promotion trên Console từng khách GCP.", False),
        ("5.2.4", "Hệ thống xuất danh sách khách hàng có credit trong tháng để gửi Ban Giám Đốc/Sale xác nhận.", False),
        ("5.2.5", "[CẦN XÁC NHẬN] Hệ thống hỗ trợ trường hợp chia credit: ví dụ 4.000 credit → 2.500 cho khách, 1.500 giữ lại.", True),
        ("5.2.6", "Hệ thống cảnh báo khi phát hiện credit chưa được phân loại/xác nhận trước khi gửi bill.", False),
        ("5.2.7", "GMP: Hệ thống xác nhận không áp dụng Credit và Gemini API cho dịch vụ GMP.", False),
    ]),
    ("5.3 Công thức Tính Cước", [
        ("5.3.1", "Hệ thống cho phép cấu hình công thức tính cước riêng cho từng khách hàng, từng hợp đồng — bao gồm: tỷ lệ discount, thuế VAT nhà thầu (VAT GG), phí dịch vụ (PDV).", False),
        ("5.3.2", "GCP: Hệ thống tính theo công thức: Thu khách = (Usage_total - Gemini) × (1 - Discount%) + Gemini + VAT GG + PDV - Credit.", False),
        ("5.3.3", "GCP: VAT GG hiện tăng từ 5,263% lên 5,88%. Một số khách đặc biệt (VD: VPBank) tính VAT GG trên số sau discount.", False),
        ("5.3.4", "GCP: PDV = (Lượng dùng trước discount + VAT GG) × Tỷ lệ PDV (5,88% hoặc 5,263% tùy phụ lục hợp đồng).", False),
        ("5.3.5", "GMP: Thu khách = Lượng dùng USD × (1 - Discount%) + PDV.", False),
        ("5.3.6", "GWS Flex: Tính từ lượng dùng trên file CSV, áp dụng SKU pricing và công thức theo ngày (đã implement trong calculateGwsFlex.js).", False),
        ("5.3.7", "Quy đổi VND: Thu khách (VND) = Thu khách (USD) × Tỷ giá. Làm tròn đến hàng đơn vị.", False),
        ("5.3.8", "Thuế GTGT đầu ra: 10% (đa phần). Có dịch vụ Không chịu thuế (KCT) như Bubble.", False),
        ("5.3.9", "[CẦN XÁC NHẬN] Hệ thống hỗ trợ cập nhật/thay đổi công thức tính khi có phụ lục hợp đồng mới.", True),
    ]),
    ("5.4 Sinh Bảng Đối soát Chi phí", [
        ("5.4.1", "Hệ thống tự động sinh bảng đối soát chi phí cho từng khách hàng, từng dịch vụ (GCP, GMP, GWS riêng biệt) dùng template XLSX trên S3.", False),
        ("5.4.2", "Hệ thống hỗ trợ khách có nhiều project: tự động cộng tổng tất cả project của cùng 1 khách (dựa trên Project number / Subaccount ID).", False),
        ("5.4.3", "Hệ thống hỗ trợ khách có nhiều billing account: tự động gộp dữ liệu từ nhiều billing ID vào cùng 1 bảng đối soát.", False),
        ("5.4.4", "Hệ thống cho phép kế toán sửa tay số liệu trên bảng đối soát trước khi gửi khách (điều chỉnh credit, lỗi rounding, phí dịch vụ).", False),
        ("5.4.5", "Hệ thống xuất bảng đối soát dưới dạng XLSX (hiện tại) và PDF (mở rộng sau).", False),
        ("5.4.6", "Bảng đối soát gửi khách đã bao gồm toàn bộ chiết khấu, phí dịch vụ — số tiền cuối cùng khách phải thanh toán.", False),
        ("5.4.7", "Hệ thống hỗ trợ đối soát tự động: so khớp số ERP gen ra với bảng tính tay của kế toán, highlight dòng lệch.", False),
        ("5.4.8", "Hệ thống xử lý chính xác việc làm tròn (đến hàng đồng VND), giải quyết lỗi rounding hiện tại của CM.", False),
    ]),
    ("5.5 Gửi Bill & Theo dõi Xác nhận Khách hàng", [
        ("5.5.1", "Hệ thống gửi email bảng đối soát (PDF/XLSX) kèm screenshot lượng dùng cho khách theo danh sách email đã đăng ký trên hợp đồng.", False),
        ("5.5.2", "Dịch vụ khác nhau gửi riêng từng email (GWS, GCP, GMP gửi lần lượt theo thời điểm có invoice).", False),
        ("5.5.3", "Khách có 02 ngày làm việc để xác nhận chi phí. Sau 02 ngày không phản hồi, hệ thống tự động gửi email nhắc nhở.", False),
        ("5.5.4", "Sau nhắc nhở 1 ngày vẫn không phản hồi, hệ thống tự động chốt số liệu và cho phép xuất hóa đơn.", False),
        ("5.5.5", "Hệ thống ghi nhận xác nhận qua 2 hình thức: phản hồi email hoặc khách thực hiện chuyển khoản thanh toán.", False),
        ("5.5.6", "Hệ thống xuất hóa đơn trực tiếp sau khi khách xác nhận, không cần bước phê duyệt trung gian.", False),
    ]),
    ("5.6 Xử lý Tranh chấp Cước (Dispute)", [
        ("5.6.1", "Khi khách phản hồi lệch cước qua email, hệ thống cho phép kế toán mở bảng đối soát, sửa đổi trực tiếp và gửi lại bản cập nhật.", False),
        ("5.6.2", "Dịch vụ nào có tranh chấp thì xử lý riêng, các dịch vụ khác vẫn chốt bình thường.", False),
        ("5.6.3", "Khi phát hiện lệch cước sau khi đã xuất hóa đơn VAT, hệ thống hỗ trợ gửi biên bản điều chỉnh, phát hành hóa đơn điều chỉnh tăng/giảm.", False),
        ("5.6.4", "[CẦN XÁC NHẬN] Hệ thống cho phép điều chỉnh vào kỳ cước tiếp theo thay vì hủy/sửa hóa đơn đã xuất.", True),
    ]),
    ("5.7 Cấu hình Hợp đồng & Khách hàng", [
        ("5.7.1", "Hệ thống cho phép cấu hình cho từng hợp đồng: tên khách hàng, dịch vụ, tỷ lệ chiết khấu, phí dịch vụ, thuế VAT GG, thời gian đối soát, thời hạn thanh toán.", False),
        ("5.7.2", "Hệ thống hỗ trợ khách thay đổi pháp nhân: giữ nguyên số liệu, thay đổi pháp nhân xuất hóa đơn.", False),
        ("5.7.3", "Hệ thống hỗ trợ khách chia nhiều pháp nhân: tính riêng, gửi riêng, xuất hóa đơn riêng cho từng pháp nhân.", False),
        ("5.7.4", "Hệ thống thông báo cho kế toán khi có thay đổi hợp đồng/phụ lục.", False),
        ("5.7.5", "Hệ thống cho phép cấu hình ngân hàng lấy tỷ giá và thời điểm lấy tỷ giá riêng cho từng hợp đồng.", False),
    ]),
    ("5.8 Quản lý Timeline Tính Cước", [
        ("5.8.1", "Hệ thống quản lý timeline tính cước theo từng dịch vụ: GWS Flex (ngày 1-3), GCP (ngày 1-6), GMP (ngày 6-10).", False),
        ("5.8.2", "Hệ thống cảnh báo khi hóa đơn hãng về chậm so với deadline cam kết gửi khách.", False),
        ("5.8.3", "Khi invoice về chậm hơn deadline hợp đồng, hệ thống hỗ trợ gửi thông báo trễ cho khách.", False),
        ("5.8.4", "Khách hủy dịch vụ giữa tháng: hệ thống chờ đến đầu tháng kế tiếp khi nhận invoice hãng.", False),
    ]),
    ("5.9 Lưu trữ & Tra cứu Lịch sử Bill", [
        ("5.9.1", "Hệ thống lưu trữ tập trung toàn bộ lịch sử bill, bảng đối soát, hóa đơn theo khách hàng/tháng/dịch vụ trên S3.", False),
        ("5.9.2", "Hệ thống cho phép Sales AM và Kế toán tra cứu bill lịch sử theo bộ lọc: khách hàng, tháng/năm, dịch vụ.", False),
        ("5.9.3", "Hệ thống lưu lịch sử tất cả lần gửi/sửa của cùng 1 bảng đối soát (versioning).", False),
        ("5.9.4", "Hệ thống xuất được file Excel tổng hợp số cuối của toàn bộ khách hàng trong tháng.", False),
    ]),
    ("5.10 Kiểm soát & Audit Trail", [
        ("5.10.1", "[CẦN XÁC NHẬN] Hệ thống ghi log mọi thay đổi trên bảng đối soát: ai sửa, lúc nào, giá trị trước/sau.", True),
        ("5.10.2", "[CẦN XÁC NHẬN] ERP cần ghi nhận hành vi sửa tay tự động (audit trail) để truy vết khi cần.", True),
    ]),
]

for group_title, reqs in requirements:
    doc.add_heading(group_title, level=2)
    for num, text, is_open in reqs:
        p = doc.add_paragraph(style='List Bullet')
        add_req(p, num, text, is_open)

# ── 6. Câu hỏi còn Mở ──
doc.add_heading("6. Câu hỏi còn Mở — Cần xác nhận với khách hàng", level=1)

open_questions = [
    ("OQ-01", "Luồng lấy cước AWS từ đầu đến cuối (portal, file, format) — chưa có phản hồi từ kế toán."),
    ("OQ-02", "Luồng lấy cước DigitalOcean từ đầu đến cuối — chưa có phản hồi."),
    ("OQ-03", "Ngoài GCP/GMP/GWS/AWS/DO còn hãng nào khác cần tính cước?"),
    ("OQ-04", "ERP mới sẽ thay thế hay tích hợp với CM? Nếu tích hợp: dữ liệu nào đẩy sang CM, dữ liệu nào CM đẩy về ERP?"),
    ("OQ-05", "ERP có cần bắt buộc ghi lý do mỗi lần kế toán chỉnh sửa số liệu bill trước khi gửi khách không?"),
    ("OQ-06", "ERP có cần lưu lịch sử version trước/sau khi sửa bill để truy vết không?"),
    ("OQ-07", "Quy trình xử lý khi kế toán sửa sai bill rồi đã gửi khách là gì?"),
    ("OQ-08", "Ngưỡng Gemini API 'quá nhỏ' ($0.07) có cần cấu hình được trên ERP hay cố định?"),
    ("OQ-09", "Cần chia sẻ file Excel tính cước mẫu cho ít nhất 1 hãng (GCP hoặc GWS) để số hóa đúng công thức trên ERP."),
    ("OQ-10", "Hiện tại ~10-15 khách GG mỗi tháng có Promotion Credit. Có danh sách cố định phân biệt Credit của hãng vs. Credit của CloudAZ không?"),
    ("OQ-11", "Quy trình onboard khách mới: ERP có cần module cấu hình công thức tính cước theo HĐ mà không cần nhập tay Excel?"),
    ("OQ-12", "SLA xử lý tranh chấp cước là bao lâu? Hiện chưa có SLA."),
    ("OQ-13", "Free tier (F2) GCP — năm thứ 2 được giảm 20%: ERP cần tự động nhận diện hay kế toán nhập tay?"),
    ("OQ-14", "Công thức GWS Flex cũ (trước 02/2024): amount = Excel.Amount * 100/80 — có cần hỗ trợ trong ERP không?"),
]

oq_table = doc.add_table(rows=1 + len(open_questions), cols=3)
oq_table.style = 'Light Shading Accent 1'
oq_table.rows[0].cells[0].text = "Mã"
oq_table.rows[0].cells[1].text = "Vấn đề cần xác nhận"
oq_table.rows[0].cells[2].text = "Trạng thái"
for i, (code, question) in enumerate(open_questions):
    oq_table.rows[i+1].cells[0].text = code
    oq_table.rows[i+1].cells[1].text = question
    oq_table.rows[i+1].cells[2].text = "Chờ xác nhận"

# ── Save ──
output_path = "C:/Users/thanh/Desktop/ERP_Cloudaz/docs/CM_docs/BRD_ERP_CloudAZ_Billing_Dispute_2026-08-20.docx"
doc.save(output_path)
print("OK: " + output_path)
